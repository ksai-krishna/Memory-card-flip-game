import os
import sys
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def set_table_borders(table, color="CCCCCC", sz="4", val="single"):
    tblPr = table._element.xpath('w:tblPr')
    if tblPr:
        borders = parse_xml(
            f'<w:tblBorders {nsdecls("w")}>'
            f'  <w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
            f'  <w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
            f'  <w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
            f'  <w:insideV w:val="none"/>'
            f'  <w:left w:val="none"/>'
            f'  <w:right w:val="none"/>'
            f'</w:tblBorders>'
        )
        tblPr[0].append(borders)

def add_heading_1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x00, 0x78, 0xD4)
    return p

def add_heading_2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x00, 0x45, 0x78)
    return p

def add_heading_3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x26, 0x26, 0x26)
    return p

def add_body_paragraph(doc, text, bold_prefix=""):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        r_pre.font.name = "Arial"
        r_pre.font.size = Pt(11)
        r_pre.font.bold = True
        r_pre.font.color.rgb = RGBColor(0x26, 0x26, 0x26)
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return p

def add_bullet_item(doc, text, bold_prefix=""):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        r_pre.font.name = "Arial"
        r_pre.font.size = Pt(10.5)
        r_pre.font.bold = True
        r_pre.font.color.rgb = RGBColor(0x26, 0x26, 0x26)
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return p

def add_figure(doc, img_path, caption_text, width=Inches(5.8)):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run()
    run.add_picture(img_path, width=width)
    
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap.paragraph_format.space_before = Pt(2)
    p_cap.paragraph_format.space_after = Pt(14)
    p_cap.paragraph_format.keep_with_next = True
    r_cap = p_cap.add_run(caption_text)
    r_cap.font.name = "Arial"
    r_cap.font.size = Pt(9.5)
    r_cap.font.italic = True
    r_cap.font.bold = True
    r_cap.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

def build_docx(output_path):
    doc = docx.Document()
    
    # Page Setup - Margins 1 inch
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
        # Header / Footer setup
        footer = section.footer
        f_p = footer.paragraphs[0]
        f_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        f_run = f_p.add_run("Cloud-Based Event Management Portal using Microsoft Azure | VTU M.Tech Mini Project")
        f_run.font.name = "Arial"
        f_run.font.size = Pt(8.5)
        f_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    # -------------------------------------------------------------
    # COVER PAGE
    # -------------------------------------------------------------
    # Logos header table
    logo_table = doc.add_table(rows=1, cols=2)
    logo_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    logo_table.autofit = False
    
    cell_l = logo_table.cell(0, 0)
    cell_r = logo_table.cell(0, 1)
    cell_l.width = Inches(3.25)
    cell_r.width = Inches(3.25)
    
    p_l = cell_l.paragraphs[0]
    p_l.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if os.path.exists("extracted_images/img-000.png"):
        p_l.add_run().add_picture("extracted_images/img-000.png", width=Inches(1.2))
        
    p_r = cell_r.paragraphs[0]
    p_r.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if os.path.exists("extracted_images/img-001.png"):
        p_r.add_run().add_picture("extracted_images/img-001.png", width=Inches(1.4))

    # Title details
    p_univ = doc.add_paragraph()
    p_univ.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_univ.paragraph_format.space_before = Pt(16)
    p_univ.paragraph_format.space_after = Pt(2)
    r_u1 = p_univ.add_run("VISVESVARAYA TECHNOLOGICAL UNIVERSITY\n")
    r_u1.bold = True
    r_u1.font.name = "Arial"
    r_u1.font.size = Pt(13)
    r_u1.font.color.rgb = RGBColor(0x00, 0x45, 0x78)
    
    r_u2 = p_univ.add_run('"JNANA SANGAMA", MACHHE, BELAGAVI-590018')
    r_u2.font.name = "Arial"
    r_u2.font.size = Pt(9.5)
    r_u2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_before = Pt(20)
    p_sub.paragraph_format.space_after = Pt(8)
    r_rep = p_sub.add_run("Mini Project Report on\n")
    r_rep.font.name = "Arial"
    r_rep.font.size = Pt(12)
    r_rep.font.italic = True

    r_title = p_sub.add_run("Cloud-Based Event Management Portal using Microsoft Azure")
    r_title.bold = True
    r_title.font.name = "Arial"
    r_title.font.size = Pt(18)
    r_title.font.color.rgb = RGBColor(0x00, 0x78, 0xD4)

    p_fulfillment = doc.add_paragraph()
    p_fulfillment.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_fulfillment.paragraph_format.space_before = Pt(10)
    p_fulfillment.paragraph_format.space_after = Pt(20)
    r_ful = p_fulfillment.add_run(
        "Submitted in partial fulfillment of the requirements for the award of the degree\n"
        "Master of Technology\n"
        "in\n"
        "Computer Science and Engineering\n"
        "of\n"
        "Visvesvaraya Technological University, Belagavi."
    )
    r_ful.font.name = "Arial"
    r_ful.font.size = Pt(10)
    r_ful.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    # Details table (Student & Guide)
    details_table = doc.add_table(rows=1, cols=2)
    details_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    details_table.autofit = False
    
    cell_st = details_table.cell(0, 0)
    cell_gd = details_table.cell(0, 1)
    cell_st.width = Inches(3.25)
    cell_gd.width = Inches(3.25)
    
    p_st = cell_st.paragraphs[0]
    p_st.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r_sb = p_st.add_run("Submitted By:\n")
    r_sb.bold = True
    r_sb.font.name = "Arial"
    r_sb.font.size = Pt(10)
    r_sb.font.color.rgb = RGBColor(0x00, 0x45, 0x78)
    
    r_st_name = p_st.add_run("PRAVEESHA PRASAD\n(1CD25SCS10)")
    r_st_name.bold = True
    r_st_name.font.name = "Arial"
    r_st_name.font.size = Pt(11)

    p_gd = cell_gd.paragraphs[0]
    p_gd.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r_gb = p_gd.add_run("Under the Guidance of:\n")
    r_gb.bold = True
    r_gb.font.name = "Arial"
    r_gb.font.size = Pt(10)
    r_gb.font.color.rgb = RGBColor(0x00, 0x45, 0x78)
    
    r_gd_name = p_gd.add_run("Arun P\nAssistant Professor\nDept. of CSE, CITech")
    r_gd_name.bold = True
    r_gd_name.font.name = "Arial"
    r_gd_name.font.size = Pt(10.5)

    # Foot college info
    p_coll = doc.add_paragraph()
    p_coll.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_coll.paragraph_format.space_before = Pt(30)
    p_coll.paragraph_format.space_after = Pt(0)
    
    r_dept = p_coll.add_run("Department of Computer Science and Engineering\n")
    r_dept.bold = True
    r_dept.font.name = "Arial"
    r_dept.font.size = Pt(11.5)
    r_dept.font.color.rgb = RGBColor(0x00, 0x45, 0x78)
    
    r_col = p_coll.add_run("CAMBRIDGE INSTITUTE OF TECHNOLOGY, BANGALORE-560036\n")
    r_col.bold = True
    r_col.font.name = "Arial"
    r_col.font.size = Pt(11)
    
    r_yr = p_coll.add_run("2025-2026")
    r_yr.font.name = "Arial"
    r_yr.font.size = Pt(11)

    doc.add_page_break()

    # -------------------------------------------------------------
    # TABLE OF CONTENTS & LIST OF FIGURES
    # -------------------------------------------------------------
    add_heading_1(doc, "Table of Contents")
    
    toc_table = doc.add_table(rows=1, cols=2)
    toc_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(toc_table, color="CCCCCC", sz="4")
    
    hdr_cells = toc_table.rows[0].cells
    hdr_cells[0].width = Inches(5.2)
    hdr_cells[1].width = Inches(1.3)
    
    set_cell_background(hdr_cells[0], "0078D4")
    set_cell_background(hdr_cells[1], "0078D4")
    set_cell_margins(hdr_cells[0], top=100, bottom=100, left=120, right=120)
    set_cell_margins(hdr_cells[1], top=100, bottom=100, left=120, right=120)

    p0 = hdr_cells[0].paragraphs[0]
    p0.paragraph_format.space_before = Pt(2)
    p0.paragraph_format.space_after = Pt(2)
    r0 = p0.add_run("Chapter & Title")
    r0.bold = True
    r0.font.name = "Arial"
    r0.font.size = Pt(10)
    r0.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    p1 = hdr_cells[1].paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p1.paragraph_format.space_before = Pt(2)
    p1.paragraph_format.space_after = Pt(2)
    r1 = p1.add_run("Page No.")
    r1.bold = True
    r1.font.name = "Arial"
    r1.font.size = Pt(10)
    r1.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    toc_items = [
        ("CHAPTER 1 – INTRODUCTION", "2", True),
        ("  1.1 Background", "2", False),
        ("  1.2 Problem Statement", "2", False),
        ("  1.3 Objectives", "3", False),
        ("  1.4 Scope of the Project", "3", False),
        ("  1.5 Research Gap", "3", False),
        ("  1.6 Proposed System", "4", False),
        ("  1.7 Contributions", "4", False),
        ("  1.8 Organization of the Report", "5", False),
        ("CHAPTER 2 – LITERATURE SURVEY", "6", True),
        ("  2.1 Introduction", "6", False),
        ("  2.2 Literature Review", "6", False),
        ("  2.3 Comparative Analysis", "8", False),
        ("  2.4 Research Gap", "8", False),
        ("CHAPTER 3 – SYSTEM ANALYSIS AND DESIGN", "9", True),
        ("  3.1 Existing System", "9", False),
        ("  3.2 Proposed System", "9", False),
        ("  3.3 System Architecture", "10", False),
        ("  3.4 Flow Diagram", "11", False),
        ("  3.5 Use Case Diagram", "11", False),
        ("CHAPTER 4 – METHODOLOGY", "12", True),
        ("CHAPTER 5 – IMPLEMENTATION", "13", True),
        ("CHAPTER 6 – RESULTS AND DISCUSSION", "14", True),
        ("CHAPTER 7 – CONCLUSION AND FUTURE SCOPE", "17", True),
    ]

    for title, pg, is_chap in toc_items:
        row_cells = toc_table.add_row().cells
        row_cells[0].width = Inches(5.2)
        row_cells[1].width = Inches(1.3)
        set_cell_margins(row_cells[0], top=60, bottom=60, left=120, right=120)
        set_cell_margins(row_cells[1], top=60, bottom=60, left=120, right=120)
        
        if is_chap:
            set_cell_background(row_cells[0], "F0F6FF")
            set_cell_background(row_cells[1], "F0F6FF")

        p_t = row_cells[0].paragraphs[0]
        p_t.paragraph_format.space_before = Pt(2)
        p_t.paragraph_format.space_after = Pt(2)
        r_t = p_t.add_run(title)
        r_t.font.name = "Arial"
        r_t.font.size = Pt(10)
        if is_chap:
            r_t.bold = True
            r_t.font.color.rgb = RGBColor(0x00, 0x45, 0x78)

        p_p = row_cells[1].paragraphs[0]
        p_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_p.paragraph_format.space_before = Pt(2)
        p_p.paragraph_format.space_after = Pt(2)
        r_p = p_p.add_run(pg)
        r_p.font.name = "Arial"
        r_p.font.size = Pt(10)
        if is_chap:
            r_p.bold = True

    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    add_heading_1(doc, "List of Figures")

    fig_table = doc.add_table(rows=1, cols=3)
    fig_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(fig_table, color="CCCCCC", sz="4")
    
    fhdr_cells = fig_table.rows[0].cells
    fhdr_cells[0].width = Inches(1.5)
    fhdr_cells[1].width = Inches(3.8)
    fhdr_cells[2].width = Inches(1.2)
    
    for c in fhdr_cells:
        set_cell_background(c, "0078D4")
        set_cell_margins(c, top=100, bottom=100, left=120, right=120)

    p0 = fhdr_cells[0].paragraphs[0]
    p0.paragraph_format.space_before = Pt(2)
    p0.paragraph_format.space_after = Pt(2)
    r0 = p0.add_run("Figure No.")
    r0.bold = True
    r0.font.name = "Arial"
    r0.font.size = Pt(10)
    r0.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    p1 = fhdr_cells[1].paragraphs[0]
    p1.paragraph_format.space_before = Pt(2)
    p1.paragraph_format.space_after = Pt(2)
    r1 = p1.add_run("Figure Title")
    r1.bold = True
    r1.font.name = "Arial"
    r1.font.size = Pt(10)
    r1.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    p2 = fhdr_cells[2].paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p2.paragraph_format.space_before = Pt(2)
    p2.paragraph_format.space_after = Pt(2)
    r2 = p2.add_run("Page No.")
    r2.bold = True
    r2.font.name = "Arial"
    r2.font.size = Pt(10)
    r2.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    fig_items = [
        ("Fig. 3.1", "System Architecture", "10"),
        ("Fig. 3.2", "Flow Diagram", "11"),
        ("Fig. 3.3", "Use Case Diagram", "11"),
        ("Fig. 5.1", "Home Page of Event Management Portal", "14"),
        ("Fig. 5.2", "About Page of Event Management Portal", "14"),
        ("Fig. 5.3", "Upcoming Events Page", "15"),
        ("Fig. 5.4", "Event Registration Page", "15"),
        ("Fig. 5.5", "Registration Success Message", "16"),
        ("Fig. 5.6", "Azure DevOps Pipeline", "16"),
    ]

    for f_no, f_title, f_pg in fig_items:
        row_cells = fig_table.add_row().cells
        row_cells[0].width = Inches(1.5)
        row_cells[1].width = Inches(3.8)
        row_cells[2].width = Inches(1.2)
        for c in row_cells:
            set_cell_margins(c, top=60, bottom=60, left=120, right=120)
            
        p_no = row_cells[0].paragraphs[0]
        p_no.paragraph_format.space_before = Pt(2)
        p_no.paragraph_format.space_after = Pt(2)
        r_n = p_no.add_run(f_no)
        r_n.font.name = "Arial"
        r_n.font.size = Pt(10)
        r_n.bold = True

        p_t = row_cells[1].paragraphs[0]
        p_t.paragraph_format.space_before = Pt(2)
        p_t.paragraph_format.space_after = Pt(2)
        r_t = p_t.add_run(f_title)
        r_t.font.name = "Arial"
        r_t.font.size = Pt(10)

        p_p = row_cells[2].paragraphs[0]
        p_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_p.paragraph_format.space_before = Pt(2)
        p_p.paragraph_format.space_after = Pt(2)
        r_p = p_p.add_run(f_pg)
        r_p.font.name = "Arial"
        r_p.font.size = Pt(10)

    doc.add_page_break()

    # -------------------------------------------------------------
    # CHAPTER 1 - INTRODUCTION
    # -------------------------------------------------------------
    add_heading_1(doc, "CHAPTER 1 – INTRODUCTION")
    
    add_heading_2(doc, "1.1 Background")
    add_body_paragraph(doc, "The rapid growth of cloud computing has transformed the way web applications are developed, deployed, and maintained. Organizations and educational institutions increasingly rely on cloud platforms to host applications because they provide high availability, scalability, reliability, and secure access from any location. Cloud platforms eliminate the need for maintaining expensive local servers while offering automatic deployment and continuous integration services.")
    add_body_paragraph(doc, "Event management is an essential activity in colleges, universities, companies, and organizations. Traditional methods of managing events often involve manual registration forms, spreadsheets, and paper-based records, making the process time-consuming and prone to errors. Participants may face difficulties in obtaining event information, registering for events, and receiving timely updates.")
    add_body_paragraph(doc, "The Cloud-Based Event Management Portal is designed to simplify the entire event management process by providing a centralized web application hosted on Microsoft Azure. The application allows users to browse upcoming events, view event schedules, register online, and receive registration confirmation through an intuitive interface.")
    add_body_paragraph(doc, "The project is developed using HTML, CSS, and JavaScript in Visual Studio Code, with the source code managed through GitHub and deployed using Azure DevOps Pipelines. The integration of GitHub with Azure enables Continuous Integration and Continuous Deployment (CI/CD), ensuring that every change made to the application is automatically deployed to the cloud. This approach demonstrates the practical implementation of cloud computing technologies in modern web application development while providing a reliable, scalable, and easily accessible event management solution.")

    add_heading_2(doc, "1.2 Problem Statement")
    add_body_paragraph(doc, "Many educational institutions and organizations continue to manage events using manual registration methods or basic web pages that lack automation and cloud deployment. These traditional systems make it difficult to maintain event information, manage participant registrations, and update content efficiently. Manual processes increase administrative workload and are susceptible to errors, duplicate registrations, and delayed communication.")
    add_body_paragraph(doc, "Additionally, locally hosted applications require continuous maintenance and do not provide the flexibility of remote accessibility or automated deployment. Every update requires manual intervention, increasing development time and reducing system reliability.")
    add_body_paragraph(doc, "To address these limitations, this project proposes a Cloud-Based Event Management Portal deployed on Microsoft Azure. The application uses Azure cloud services and GitHub integration to provide automated deployment, centralized management, secure hosting, and easy accessibility.")

    add_heading_2(doc, "1.3 Objectives")
    add_bullet_item(doc, "To develop a cloud-based Event Management Portal using Microsoft Azure.")
    add_bullet_item(doc, "To provide a centralized platform for displaying event information.")
    add_bullet_item(doc, "To allow users to register for events through an online registration form.")
    add_bullet_item(doc, "To display upcoming events and event schedules in an organized manner.")
    add_bullet_item(doc, "To automate the deployment process using GitHub and Azure DevOps Pipelines.")
    add_bullet_item(doc, "To demonstrate Continuous Integration and Continuous Deployment (CI/CD) practices.")
    add_bullet_item(doc, "To enhance user accessibility by making the application available online.")

    add_heading_2(doc, "1.4 Scope of the Project")
    add_body_paragraph(doc, "The scope of this project includes the design, development, and cloud deployment of a web-based event management application. The project focuses on creating a user-friendly frontend using HTML, CSS, and JavaScript, while managing the source code using GitHub and deploying the application on Microsoft Azure using Azure DevOps Pipelines.")
    add_body_paragraph(doc, "The application is suitable for educational institutions, colleges, and small organizations to manage events, display schedules, and collect participant registrations. Future enhancements may include backend databases, user authentication, automated email notifications, and mobile app integration.")

    add_heading_2(doc, "1.5 Research Gap")
    add_body_paragraph(doc, "While several commercial event management tools exist, they are often complex, costly, and do not provide clear insight into cloud deployment mechanisms for educational purposes. Most academic mini-projects focus solely on local web development without incorporating production-grade cloud deployment pipelines. This project bridges that gap by demonstrating a full CI/CD deployment workflow using Azure DevOps and GitHub.")

    add_heading_2(doc, "1.6 Proposed System")
    add_body_paragraph(doc, "The proposed system provides a responsive, web-based portal hosted on Microsoft Azure App Service. Users can view event schedules, detailed descriptions, and submit registrations through client-side validated forms. Developers commit updates to GitHub, which automatically triggers Azure DevOps CI/CD pipelines to build and deploy updates seamlessly.")

    add_heading_2(doc, "1.7 Contributions")
    add_bullet_item(doc, "Designed and implemented a responsive web interface for event management.")
    add_bullet_item(doc, "Integrated GitHub version control with Azure DevOps Pipelines for automated CI/CD.")
    add_bullet_item(doc, "Successfully deployed a live application on Microsoft Azure App Service.")
    add_bullet_item(doc, "Demonstrated practical implementation of cloud-hosted web architecture for educational institutions.")

    add_heading_2(doc, "1.8 Organization of the Report")
    add_bullet_item(doc, "Chapter 1 presents the introduction, problem statement, objectives, scope, and proposed system.")
    add_bullet_item(doc, "Chapter 2 presents the literature survey on cloud platforms, Azure DevOps, and web-based systems.")
    add_bullet_item(doc, "Chapter 3 discusses system analysis and design, including architecture, flow diagram, and use case diagram.")
    add_bullet_item(doc, "Chapter 4 explains the methodology adopted for user interface, event registration, and automated deployment.")
    add_bullet_item(doc, "Chapter 5 describes implementation details, technologies used, and application modules.")
    add_bullet_item(doc, "Chapter 6 presents application results, user interface screenshots, and pipeline execution.")
    add_bullet_item(doc, "Chapter 7 concludes the report and discusses potential future enhancements.")

    doc.add_page_break()

    # -------------------------------------------------------------
    # CHAPTER 2 - LITERATURE SURVEY
    # -------------------------------------------------------------
    add_heading_1(doc, "CHAPTER 2 – LITERATURE SURVEY")
    
    add_heading_2(doc, "2.1 Introduction")
    add_body_paragraph(doc, "Cloud computing has significantly transformed the development and deployment of modern web applications by providing scalable, reliable, and cost-effective infrastructure. Organizations and educational institutions increasingly use cloud platforms such as Microsoft Azure to host applications, ensuring high availability, secure access, and simplified maintenance. Along with cloud hosting, version control systems like GitHub and DevOps practices have become essential for collaborative software development and automated deployment.")
    add_body_paragraph(doc, "Event management systems are widely used to organize conferences, workshops, seminars, technical fests, and cultural programs. Traditional event management methods often rely on manual registration and static web pages, which are difficult to maintain and update. Cloud-based event management portals address these challenges by offering centralized access to event information, online registration, and automated deployment through cloud platforms.")

    add_heading_2(doc, "2.2 Literature Review")
    
    add_heading_3(doc, "Paper 1: Microsoft Azure App Service: A Scalable Platform for Cloud Web Applications (2023)")
    add_body_paragraph(doc, "Authors: Microsoft Azure Documentation\nThis work describes Microsoft Azure App Service as a cloud platform that enables developers to build, deploy, and host web applications without managing physical servers. It supports automatic scaling, secure deployment, high availability, and seamless integration with Azure DevOps and GitHub.")
    add_bullet_item(doc, "Provides secure and reliable cloud hosting.", "Advantages: ")
    add_bullet_item(doc, "Supports automatic deployment and scaling.", "Advantages: ")
    add_bullet_item(doc, "Easily integrates with GitHub and Azure DevOps.", "Advantages: ")
    add_bullet_item(doc, "Requires internet connectivity; initial Azure configuration may be challenging for beginners.", "Limitations: ")

    add_heading_3(doc, "Paper 2: GitHub and CI/CD for Modern Web Application Development (2022)")
    add_body_paragraph(doc, "Authors: GitHub Documentation\nThis study explains the importance of GitHub as a distributed version control platform that enables collaborative software development. By integrating GitHub with Continuous Integration and Continuous Deployment (CI/CD) pipelines, developers can automatically build, test, and deploy web applications whenever code changes are committed.")
    add_bullet_item(doc, "Efficient version control and collaboration; automated application deployment.", "Advantages: ")
    add_bullet_item(doc, "Requires understanding of Git commands; pipeline configuration can be complex.", "Limitations: ")

    add_heading_3(doc, "Paper 3: Web-Based Event Management Systems for Educational Institutions (2024)")
    add_body_paragraph(doc, "Authors: Various Researchers\nThis study focuses on web-based event management systems that simplify event organization through online registration and centralized event information. The research highlights that digital event management improves user experience, reduces administrative workload, and enables participants to access event schedules from any location.")
    add_bullet_item(doc, "Simplifies event registration and improves participant accessibility.", "Advantages: ")
    add_bullet_item(doc, "Requires regular updates; basic systems often lack cloud deployment automation.", "Limitations: ")

    add_heading_2(doc, "2.3 Comparative Analysis")
    
    comp_table = doc.add_table(rows=1, cols=4)
    comp_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(comp_table, color="CCCCCC", sz="4")
    
    chdr_cells = comp_table.rows[0].cells
    chdr_cells[0].width = Inches(1.6)
    chdr_cells[1].width = Inches(1.5)
    chdr_cells[2].width = Inches(1.7)
    chdr_cells[3].width = Inches(1.7)
    
    for c in chdr_cells:
        set_cell_background(c, "0078D4")
        set_cell_margins(c, top=100, bottom=100, left=100, right=100)

    titles = ["Paper", "Technique", "Advantages", "Limitations"]
    for i, t in enumerate(titles):
        p = chdr_cells[i].paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(t)
        r.bold = True
        r.font.name = "Arial"
        r.font.size = Pt(9.5)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    rows_data = [
        ("Microsoft Azure App Service", "Cloud Web Hosting", "Scalable, secure, and reliable application hosting", "Requires internet connectivity; learning curve for Azure"),
        ("GitHub with CI/CD", "Version Control & Automated Deployment", "Efficient collaboration and automatic deployment", "Requires Git knowledge and pipeline setup"),
        ("Web-Based Event Management System", "Online Event Management", "Easy registration and centralized event info", "Requires continuous maintenance and manual updates"),
    ]

    for r_data in rows_data:
        row_cells = comp_table.add_row().cells
        for i, val in enumerate(r_data):
            row_cells[i].width = [Inches(1.6), Inches(1.5), Inches(1.7), Inches(1.7)][i]
            set_cell_margins(row_cells[i], top=60, bottom=60, left=100, right=100)
            p = row_cells[i].paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(val)
            r.font.name = "Arial"
            r.font.size = Pt(9.5)

    add_heading_2(doc, "2.4 Research Gap")
    add_body_paragraph(doc, "The literature survey indicates that existing event management systems primarily focus on online registration and basic website functionality, while cloud hosting platforms mainly emphasize scalable application deployment. Similarly, GitHub provides efficient version control, but many educational projects do not demonstrate the complete integration of GitHub, Azure DevOps, and Microsoft Azure within a single application.")

    doc.add_page_break()

    # -------------------------------------------------------------
    # CHAPTER 3 - SYSTEM ANALYSIS AND DESIGN
    # -------------------------------------------------------------
    add_heading_1(doc, "CHAPTER 3 – SYSTEM ANALYSIS AND DESIGN")
    
    add_heading_2(doc, "3.1 Existing System")
    add_body_paragraph(doc, "Traditional event management systems are primarily based on manual registration methods, spreadsheets, or simple static websites. Event organizers collect participant details manually, making record-keeping labor-intensive and prone to errors.")
    add_bullet_item(doc, "Manual registration and event management.")
    add_bullet_item(doc, "Time-consuming process for updating event details.")
    add_bullet_item(doc, "Limited accessibility and lack of cloud hosting.")
    add_bullet_item(doc, "No automated deployment mechanism.")

    add_heading_2(doc, "3.2 Proposed System")
    add_body_paragraph(doc, "The proposed Cloud-Based Event Management Portal provides a responsive web application built with HTML, CSS, and JavaScript. It features automated deployment via Azure DevOps Pipelines and GitHub integration, hosting the portal reliably on Microsoft Azure App Service.")
    add_bullet_item(doc, "Responsive and user-friendly web interface.")
    add_bullet_item(doc, "Cloud-based deployment using Microsoft Azure App Service.")
    add_bullet_item(doc, "Automated deployment through Azure DevOps CI/CD Pipelines.")
    add_bullet_item(doc, "Version control using GitHub repository.")
    add_bullet_item(doc, "Streamlined online event registration.")

    add_heading_2(doc, "3.3 System Architecture")
    add_body_paragraph(doc, "The system architecture consists of a presentation layer, application layer, version control, CI/CD pipeline, and cloud hosting platform:")
    add_bullet_item(doc, "User accesses portal via modern web browser on desktop, tablet, or smartphone.")
    add_bullet_item(doc, "Frontend presents event schedules, details, and registration forms.")
    add_bullet_item(doc, "JavaScript validates input client-side before processing.")
    add_bullet_item(doc, "Source code maintained in GitHub; code pushes trigger Azure DevOps Pipelines.")
    add_bullet_item(doc, "Azure DevOps builds and deploys latest code directly to Azure App Service.")

    add_heading_2(doc, "3.4 Flow Diagram")
    if os.path.exists("extracted_images/img-002.png"):
        add_figure(doc, "extracted_images/img-002.png", "Fig. 3.2: Flow Diagram of Cloud-Based Event Management Portal", width=Inches(4.5))

    add_heading_2(doc, "3.5 Use Case Diagram")
    if os.path.exists("extracted_images/img-003.png"):
        add_figure(doc, "extracted_images/img-003.png", "Fig. 3.3: Use Case Diagram for Users and Developers", width=Inches(5.5))

    doc.add_page_break()

    # -------------------------------------------------------------
    # CHAPTER 4 - METHODOLOGY
    # -------------------------------------------------------------
    add_heading_1(doc, "CHAPTER 4 – METHODOLOGY")
    add_body_paragraph(doc, "The proposed Cloud-Based Event Management Portal follows a structured web engineering methodology:")
    add_bullet_item(doc, "Frontend Design: Developing modular HTML5 pages styled with CSS3 and enhanced with client-side JavaScript.")
    add_bullet_item(doc, "Modular Components: Designing dedicated views for Home, About, Upcoming Events, Event Schedule, Event Registration, Gallery, and Contact.")
    add_bullet_item(doc, "Input Validation: Utilizing JavaScript regex and field validation to verify email format, phone numbers, and mandatory fields prior to submission.")
    add_bullet_item(doc, "Version Control Workflow: Managing codebase with Git and pushing changes to GitHub repository.")
    add_bullet_item(doc, "CI/CD Pipeline Setup: Configuring Azure DevOps Pipelines connected to GitHub repo for automated build, artifact creation, and Azure App Service deployment.")

    doc.add_page_break()

    # -------------------------------------------------------------
    # CHAPTER 5 - IMPLEMENTATION
    # -------------------------------------------------------------
    add_heading_1(doc, "CHAPTER 5 – IMPLEMENTATION")
    add_body_paragraph(doc, "The implementation was carried out using standard web technologies and cloud service integrations:")
    add_bullet_item(doc, "HTML5: Structuring semantic content across application pages.")
    add_bullet_item(doc, "CSS3: Styling modern visual components, grid layouts, and responsive flex containers.")
    add_bullet_item(doc, "JavaScript (ES6): Handling user interactions, dynamic form validation, and confirmation overlays.")
    add_bullet_item(doc, "GitHub: Hosting repository and handling source code commits.")
    add_bullet_item(doc, "Azure DevOps: Managing build pipelines and release triggers.")
    add_bullet_item(doc, "Microsoft Azure App Service: Hosting live production web application with continuous availability.")

    doc.add_page_break()

    # -------------------------------------------------------------
    # CHAPTER 6 - RESULTS AND DISCUSSION
    # -------------------------------------------------------------
    add_heading_1(doc, "CHAPTER 6 – RESULTS AND DISCUSSION")
    add_body_paragraph(doc, "The proposed Cloud-Based Event Management Portal using Microsoft Azure was successfully designed, developed, and tested. The application provides an intuitive, high-performance web portal for browsing events and managing participant registrations.")

    add_heading_2(doc, "6.1 Home Page & Navigation")
    add_body_paragraph(doc, "The Home page presents a welcome banner, highlights upcoming college events, and offers navigation across all modules.")
    if os.path.exists("extracted_images/img-004.png"):
        add_figure(doc, "extracted_images/img-004.png", "Fig. 5.1: Home Page of Event Management Portal", width=Inches(5.5))

    add_heading_2(doc, "6.2 About Page")
    add_body_paragraph(doc, "The About page outlines portal objectives, college department highlights, and event guidelines.")
    if os.path.exists("extracted_images/img-005.png"):
        add_figure(doc, "extracted_images/img-005.png", "Fig. 5.2: About Page of Event Management Portal", width=Inches(5.2))

    add_heading_2(doc, "6.3 Upcoming Events Page")
    add_body_paragraph(doc, "Displays structured event cards with dates, venues, categories, and registration call-to-action buttons.")
    if os.path.exists("extracted_images/img-006.png"):
        add_figure(doc, "extracted_images/img-006.png", "Fig. 5.3: Upcoming Events Page", width=Inches(5.2))

    add_heading_2(doc, "6.4 Event Registration & Confirmation")
    add_body_paragraph(doc, "Users complete registration forms with real-time field validation. Upon successful submission, a confirmation alert is displayed.")
    if os.path.exists("extracted_images/img-007.png"):
        add_figure(doc, "extracted_images/img-007.png", "Fig. 5.4: Event Registration Page", width=Inches(5.2))
    if os.path.exists("extracted_images/img-008.png"):
        add_figure(doc, "extracted_images/img-008.png", "Fig. 5.5: Registration Success Message", width=Inches(5.2))

    add_heading_2(doc, "6.5 Azure DevOps CI/CD Pipeline Execution")
    add_body_paragraph(doc, "Demonstrates automated build and deployment execution upon code push to GitHub.")
    if os.path.exists("extracted_images/img-009.png"):
        add_figure(doc, "extracted_images/img-009.png", "Fig. 5.6: Azure DevOps Pipeline Build & Release", width=Inches(5.5))

    doc.add_page_break()

    # -------------------------------------------------------------
    # CHAPTER 7 - CONCLUSION AND FUTURE SCOPE
    # -------------------------------------------------------------
    add_heading_1(doc, "CHAPTER 7 – CONCLUSION AND FUTURE SCOPE")
    add_heading_2(doc, "7.1 Conclusion")
    add_body_paragraph(doc, "The Cloud-Based Event Management Portal using Microsoft Azure and GitHub was successfully designed, developed, and deployed as a responsive web application that simplifies the process of managing events and participant registrations. The application provides users with an intuitive interface to browse upcoming events, view schedules, and register online while demonstrating the practical implementation of cloud computing technologies. The integration of GitHub with Azure DevOps Pipelines enables Continuous Integration and Continuous Deployment (CI/CD), ensuring that every code update is automatically built and deployed to Microsoft Azure App Service. The project achieves its objectives by providing a reliable, scalable, and user-friendly cloud-based event management solution.")

    add_heading_2(doc, "7.2 Future Scope")
    add_body_paragraph(doc, "In the future, the Event Management Portal can be enhanced by:")
    add_bullet_item(doc, "Integrating cloud databases such as Azure SQL Database to store participant and event records securely.")
    add_bullet_item(doc, "Implementing role-based user authentication and dedicated admin dashboards.")
    add_bullet_item(doc, "Adding automated email/SMS confirmation notifications upon registration.")
    add_bullet_item(doc, "Integrating online payment gateways for paid workshops or conferences.")
    add_bullet_item(doc, "Generating QR-code-based digital event passes for instant check-in at venues.")
    add_bullet_item(doc, "Embedding Power BI analytics dashboards for event attendance and demographic reporting.")

    # Save document
    doc.save(output_path)
    print(f"Successfully generated DOCX report at: {output_path}")

if __name__ == "__main__":
    out_file = "/home/sai/Downloads/Memory-card-flip-game/cc_final_report.docx"
    build_docx(out_file)
